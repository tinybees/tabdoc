#!/usr/bin/env python3
# coding=utf-8

"""
@author: guoyanfeng
@software: PyCharm
@time: 19-2-11 下午6:14
"""
from collections import MutableSequence
from itertools import zip_longest
from typing import Any, List, Optional, Tuple, Union

from docx import Document, document, table
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
# noinspection PyProtectedMember
from docx.table import _Cell, _Row
from docx.text.run import Run
from path import Path

__all__ = ("WordWriter", "ValueAttr")


class ValueAttr(object):
    """
    单元格值属性
    """

    def __init__(self, value: Any, bgcolor: Optional[str] = None, halignment: str = "center", is_bold: bool = False):
        """
            单元格值属性
        Args:
            value: 表格值
            bgcolor: 表格背景颜色
            halignment: 单元格水平对齐方式
            is_bold: 是否加粗
        """
        self.value: Any = value
        self.bgcolor: Optional[str] = bgcolor
        halignment_lower = halignment.lower()
        if halignment_lower == "center":
            self.halignment: str = WD_TABLE_ALIGNMENT.CENTER
        elif halignment_lower == "left":
            self.halignment = WD_TABLE_ALIGNMENT.LEFT
        elif halignment_lower == "right":
            self.halignment = WD_TABLE_ALIGNMENT.RIGHT
        else:
            ValueError("halignment值错误，应为center、left和right之一")
        self.is_bold: bool = is_bold


class WordWriter(object):
    """
    word writer
    """

    def __init__(self, word_name, word_path=None):
        """
            word writer
        Args:
            word_name: word 名称
            word_path: word path
        """
        self.word_name = f"{word_name}.docx"
        self.word_path = word_path
        template_path = Path(__file__).dirname().joinpath("templates/template.docx").abspath()
        self.document: document.Document = Document(template_path)

    def __enter__(self):
        """

        Args:

        Returns:

        """
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        """

        Args:

        Returns:

        """
        self.save()

    # noinspection DuplicatedCode
    @staticmethod
    def _reduce_datetimes(row):
        """Receives a row, converts datetimes to strings."""

        row = list(row)

        for i, val in enumerate(row):
            if hasattr(val, "strftime"):
                row[i] = val.strftime("%Y-%m-%d %H:%M:%S")
            elif hasattr(val, 'isoformat'):
                row[i] = val.isoformat()
        return tuple(row)

    def add_table(self, header_name: str, header_data: List[List[Union[ValueAttr, str]]],
                  table_data: List[List[Union[ValueAttr, str]]],
                  merge_cells: List[Tuple[Tuple[int, int], Tuple[int, int]]] = None,
                  unit=None, body_fontsize=10):
        """
        为Word文档中添加表格
        Args:
            header_name: 表格的表头文字
            header_data: 表格的表头数据，可能有多个
            table_data: 表格的body数据，可能有多个
            merge_cells: 要合并的单元格
            unit: 表格数据的单位
            body_fontsize: body的字号大小
        Returns:

        """
        if not isinstance(header_data, MutableSequence):
            raise ValueError("header data值类型错误,请检查")
        if not isinstance(table_data, MutableSequence):
            raise ValueError("table data值类型错误,请检查")
        if not isinstance(merge_cells, MutableSequence):
            raise ValueError("merge cells值类型错误,请检查")
        for value in header_data:
            if not isinstance(value, MutableSequence):
                raise ValueError("header data值类型错误,请检查")
        for value in table_data:
            if not isinstance(value, MutableSequence):
                raise ValueError("table data值类型错误,请检查")

        p = self.document.add_paragraph(style="p-first-line-not-indent-center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        bold_run: Run = p.add_run(header_name)
        bold_run.font.size = Pt(12)
        bold_run.font.bold = True

        if unit:
            unit_p = self.document.add_paragraph()
            unit_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            unit_body = unit_p.add_run(f"单位：{unit}")
            unit_body.font.size = Pt(10.5)

        # 取倒数第一个header检查列的数量
        header_row, table_cols = len(header_data), len(header_data[-1])
        # analysis-data需要模板中指定，指定的方式要简单很多
        table_: table.Table = self.document.add_table(header_row, table_cols, "analysis-data")
        for value in merge_cells:
            if not isinstance(value, MutableSequence) and len(value) != 2:
                raise ValueError("merge cells值类型错误,请检查")
            for row_index, col_index in value:
                if row_index > header_row or col_index > table_cols:
                    raise ValueError("merge cells值错误,请检查")

        # 合并单元格
        if merge_cells is not None:
            for cell, other_cell in merge_cells:
                table_.cell(*cell).merge(table_.cell(*other_cell))
        # 设置表居中和自适应
        table_.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_.autofit = True
        # 添加表头，表头有可能有多行
        for index, header in enumerate(header_data):
            row = table_.rows[index]
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            for col_index, cell_value in enumerate(header):
                self._add_cell_value(row.cells[col_index], cell_value, body_fontsize)
        # 添加表体
        for row_data in table_data:
            row = table_.add_row()
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            for col_index, cell_value in enumerate(row_data):
                self._add_cell_value(row.cells[col_index], cell_value, body_fontsize)

        self.document.add_paragraph()  # 增加一个空行的段落

    def add_table2(self, header_name: str, rows_cols: Tuple[int, int],
                   table_data: List[List[Union[ValueAttr, str]]],
                   merge_cells: List[Tuple[Tuple[int, int], Tuple[int, int]]] = None,
                   unit=None, body_fontsize=10):
        """
        为Word文档中添加表格
        Args:
            header_name: 表格的表头文字
            rows_cols: 报表的行列
            table_data: 表格的body数据，可能有多个
            merge_cells: 要合并的单元格
            unit: 表格数据的单位
            body_fontsize: body的字号大小
        Returns:

        """
        if not isinstance(table_data, MutableSequence):
            raise ValueError("table data值类型错误,请检查")
        if not isinstance(merge_cells, MutableSequence):
            raise ValueError("merge cells值类型错误,请检查")
        for value in table_data:
            if not isinstance(value, MutableSequence):
                raise ValueError("table data值类型错误,请检查")

        p = self.document.add_paragraph(style="p-first-line-not-indent-center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        bold_run: Run = p.add_run(header_name)
        bold_run.font.size = Pt(12)
        bold_run.font.bold = True

        if unit:
            unit_p = self.document.add_paragraph()
            unit_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            unit_body = unit_p.add_run(f"单位：{unit}")
            unit_body.font.size = Pt(10.5)

        # analysis-data需要模板中指定，指定的方式要简单很多
        table_: table.Table = self.document.add_table(*rows_cols, "analysis-data")
        for value in merge_cells:
            if not isinstance(value, MutableSequence) and len(value) != 2:
                raise ValueError("merge cells值类型错误,请检查")
            for row_index, col_index in value:
                if row_index > rows_cols[0] or col_index > rows_cols[1]:
                    raise ValueError("merge cells值错误,请检查")

        # 合并单元格
        if merge_cells:
            for cell, other_cell in merge_cells:
                table_.cell(*cell).merge(table_.cell(*other_cell))
        # 设置表居中和自适应
        table_.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_.autofit = True
        # 添加表头，添加表体
        for index, row_data in enumerate(table_data):
            row = table_.rows[index]
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            for col_index, cell_value in enumerate(row_data):
                self._add_cell_value(row.cells[col_index], cell_value, body_fontsize)

        self.document.add_paragraph()  # 增加一个空行的段落

    def _add_cell_value(self, cell: _Cell, cell_value: Union[ValueAttr, str], fontsize: int):
        """
        添加单元格的值和样式
        Args:
            cell: 表格中的单元格
            cell_value: 单元格值
            fontsize: 字号大小
        Returns:

        """
        if isinstance(cell_value, ValueAttr):
            # cell.text = str(cell_value.value)
            # 因为样式中的字号无效，只能手动设置字号
            cell_body = cell.paragraphs[0].add_run(str(cell_value.value))
            cell_body.font.size = Pt(fontsize)
            if cell_value.is_bold:
                cell_body.font.bold = True

            if cell_value.bgcolor:
                self.set_cell_bgcolor(cell, cell_value.bgcolor)
            if cell_value.halignment:
                self.set_cell_halignment(cell, cell_value.halignment)  # 水平对齐
        else:
            # cell.text = str(cell_value)
            # 因为样式中的字号无效，只能手动设置字号
            cell_body = cell.paragraphs[0].add_run(str(cell_value))
            cell_body.font.size = Pt(fontsize)
            self.set_cell_halignment(cell, WD_TABLE_ALIGNMENT.CENTER)  # 水平居中
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直居中

    @staticmethod
    def set_cell_halignment(cell: _Cell, alignment: str = WD_TABLE_ALIGNMENT.CENTER):
        """
        设置单元格的水平对齐方式

        默认为水平居中
        Args:
            cell: 单元格
            alignment: 对齐方式默认居中
        Returns:

        """
        cell.paragraphs[0].paragraph_format.alignment = alignment

    @staticmethod
    def set_cell_bgcolor(cell: _Cell, rgbcolor: str):
        """
        设置单元格的背景颜色
        Args:
            cell: 单元格
            rgbcolor: 要设置的RGB颜色
        Returns:

        """
        bgcolor_style = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), rgbcolor))
        # noinspection PyProtectedMember
        cell._tc.get_or_add_tcPr().append(bgcolor_style)

    def set_row_bgcolor(self, row: _Row, rgbcolor: str):
        """
        设整行的背景颜色
        Args:
            row: 要设置的整行
            rgbcolor: 要设置的RGB颜色
        Returns:

        """
        for cell in row.cells:
            self.set_cell_bgcolor(cell, rgbcolor)

    @staticmethod
    def set_cell_borders(cell: _Cell, **kwargs):
        """
        Set cell`s border
        Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        # noinspection PyProtectedMember
        tc_pr = cell._tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tc_borders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tc_borders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def add_paragraph(self, bold_text: List[str] = None, *, other_text: List[str] = None):
        """
        为Word文档中添加段落
        Args:
            bold_text: 段落开头需要加粗的文本
            other_text: 段落正文的内容
        Returns:

        """
        if not isinstance(bold_text, MutableSequence):
            bold_text = [bold_text]
        if not isinstance(other_text, MutableSequence):
            other_text = [other_text]

        if bold_text and other_text:
            p = self.document.add_paragraph(style="p-first-line-indent")
            for bold_text_, other_text_ in zip_longest(bold_text, other_text):
                # 需要先添加other中的内容，因为加粗的文本前面有可能还有其他文本
                if other_text_:
                    p.add_run(other_text_)
                if bold_text_:
                    bold_run = p.add_run(bold_text_)
                    bold_run.font.size = Pt(12)
                    bold_run.font.bold = True
                    bold_run.font.color.rgb = RGBColor(255, 0, 0)
        elif bold_text and not other_text:
            p = self.document.add_paragraph(style="p-first-line-indent")
            bold_run = p.add_run("".join(bold_text))
            bold_run.font.size = Pt(12)
            bold_run.font.bold = True
            bold_run.font.color.rgb = RGBColor(255, 0, 0)
        elif not bold_text and other_text:
            self.document.add_paragraph("".join(other_text), style="p-first-line-indent")
        else:
            self.document.add_paragraph(style="p-first-line-not-indent")

    def add_heading(self, head_text: str = None, *, level: int = 1):
        """
        为Word文档中添加标题
        Args:
            head_text: 段落标题的内容
            level: 段落的级别， 共一到六级别
        Returns:

        """
        if level < 1 or level > 6:
            raise ValueError("level必须在1和6之间。")
        self.document.add_paragraph(head_text, style=f"heading{level}")

    def add_picture(self, image_path=None, image_text: str = None):
        """
        为Word文档中添加图片
        Args:
            image_path: 图片在本地的路径
            image_text: 针对图片的说明文字
        Returns:

        """
        p = self.document.add_paragraph(style="p-first-line-not-indent-center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        if image_path:
            run.add_picture(image_path, width=Inches(5.8))
        bold_run = p.add_run(f"\n{image_text}")
        bold_run.font.size = Pt(12)
        bold_run.font.bold = True

    def save(self, ):
        """
        保存工作簿
        Args:

        Returns:

        """
        if self.word_path is None:
            file_path = self.word_name
        else:
            file_path = Path(self.word_path).joinpath(self.word_name).abspath()

        self.document.save(file_path)
